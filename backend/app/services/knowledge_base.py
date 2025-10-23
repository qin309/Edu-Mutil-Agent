"""
Knowledge base service using LightRAG with multi-space support
"""
import os
import sys
from typing import List, Dict, Any, Optional
from pathlib import Path
import json
import asyncio

# Import LightRAG (assuming it's installed via pip)
try:
    from lightrag import LightRAG, QueryParam
    from lightrag.utils import EmbeddingFunc
    from lightrag.kg.shared_storage import initialize_pipeline_status
    import os

    async def siliconflow_embed(texts: List[str], model: str = "Qwen/Qwen3-Embedding-4B") -> List[List[float]]:
        """SiliconFlow async embedding function for LightRAG with configurable model and retry logic"""
        import httpx

        max_retries = 3
        retry_delay = 2  # seconds
        max_text_length = 2000  # Maximum characters per text chunk
        max_batch_size = 4  # Maximum texts per request to avoid 413 errors

        for attempt in range(max_retries):
            try:
                # Get API key from environment or use default
                api_key = os.getenv("SILICONFLOW_API_KEY", "sk-mczrhquoybaxwicxgsfdaereipgbapmtiersywbzkorjogla")

                print(f"[Embedding] Processing {len(texts)} text(s) with model '{model}', first text length: {len(texts[0]) if texts else 0} (attempt {attempt + 1}/{max_retries})")

                # Split texts into smaller chunks if they're too long
                processed_texts = []
                for text in texts:
                    if len(text) > max_text_length:
                        # Split long text into smaller chunks
                        chunks = [text[i:i+max_text_length] for i in range(0, len(text), max_text_length)]
                        processed_texts.extend(chunks)
                        print(f"[Embedding] Split long text ({len(text)} chars) into {len(chunks)} chunks")
                    else:
                        processed_texts.append(text)

                # Process texts in smaller batches to avoid 413 errors
                all_embeddings = []
                for i in range(0, len(processed_texts), max_batch_size):
                    batch = processed_texts[i:i+max_batch_size]
                    print(f"[Embedding] Processing batch {i//max_batch_size + 1}/{(len(processed_texts) + max_batch_size - 1)//max_batch_size} with {len(batch)} texts")
                    
                    async with httpx.AsyncClient(timeout=180.0) as client:
                        response = await client.post(
                            "https://api.siliconflow.cn/v1/embeddings",
                            headers={
                                "Authorization": f"Bearer {api_key}",
                                "Content-Type": "application/json",
                            },
                            json={
                                "model": model,
                                "input": batch,
                                "encoding_format": "float"
                            }
                        )
                        response.raise_for_status()
                        result = response.json()
                        batch_embeddings = [item["embedding"] for item in result["data"]]
                        all_embeddings.extend(batch_embeddings)
                
                # If we split texts, we need to average the embeddings for each original text
                if len(all_embeddings) > len(texts):
                    # Reconstruct embeddings for original texts
                    final_embeddings = []
                    embedding_idx = 0
                    for text in texts:
                        if len(text) > max_text_length:
                            # Average embeddings for chunks of this text
                            num_chunks = (len(text) + max_text_length - 1) // max_text_length
                            chunk_embeddings = all_embeddings[embedding_idx:embedding_idx + num_chunks]
                            # Average the embeddings
                            import numpy as np
                            avg_embedding = np.mean(chunk_embeddings, axis=0).tolist()
                            final_embeddings.append(avg_embedding)
                            embedding_idx += num_chunks
                        else:
                            final_embeddings.append(all_embeddings[embedding_idx])
                            embedding_idx += 1
                    all_embeddings = final_embeddings
                
                print(f"[Embedding] Successfully generated {len(all_embeddings)} embeddings, dimension: {len(all_embeddings[0]) if all_embeddings else 0}")
                return all_embeddings
                
            except Exception as e:
                print(f"[Embedding] SiliconFlow embedding error (attempt {attempt + 1}/{max_retries}): {e}")
                if attempt < max_retries - 1:
                    print(f"[Embedding] Retrying after {retry_delay} seconds...")
                    await asyncio.sleep(retry_delay)
                    retry_delay *= 2  # Exponential backoff
                else:
                    print(f"[Embedding] All retry attempts failed")
                    import traceback
                    traceback.print_exc()
                    # Fallback: return zero embeddings (adjust dimension based on model)
                    # Qwen3-Embedding-4B: 2560 dimensions, BAAI/bge-large-zh-v1.5: 1024 dimensions
                    embedding_dim = 2560 if "Qwen3-Embedding" in model else 1024
                    print(f"[Embedding] Using fallback zero embeddings with dimension {embedding_dim} for model {model}")
                    return [[0.0] * embedding_dim] * len(texts)

    # Function to set embedding dimension based on model
    def set_embedding_dim(model: str):
        if "Qwen3-Embedding" in model:
            siliconflow_embed.embedding_dim = 2560  # Qwen3-Embedding-4B dimension
        elif "bge-large" in model:
            siliconflow_embed.embedding_dim = 1024  # BAAI/bge-large-zh-v1.5 dimension
        elif "bge-m3" in model:
            siliconflow_embed.embedding_dim = 1024
        elif "bge-base" in model:
            siliconflow_embed.embedding_dim = 768
        else:
            # Default to Qwen3-Embedding-4B dimension based on user preference
            siliconflow_embed.embedding_dim = 2560
        print(f"[Embedding] Set embedding dimension to {siliconflow_embed.embedding_dim} for model {model}")
    
    # Set default embedding dimension for Qwen3-Embedding-4B (user's preferred model)
    siliconflow_embed.embedding_dim = 2560
    siliconflow_embed.set_dim = set_embedding_dim

    # Create a simple custom tokenizer for Qwen models (avoiding tiktoken issues)
    class SimpleTokenizer:
        """Simple tokenizer that uses basic character-based tokenization"""
        def __init__(self, model_name: str = "qwen"):
            self.model_name = model_name

        def encode(self, text: str) -> List[int]:
            """Simple encoding: treat each character as a token"""
            return [ord(c) for c in text]

        def decode(self, tokens: List[int]) -> str:
            """Simple decoding: convert token IDs back to characters"""
            try:
                return ''.join([chr(t) for t in tokens if 0 <= t <= 1114111])
            except (ValueError, OverflowError):
                return ''

        def __call__(self, text: str) -> List[int]:
            """Allow tokenizer to be called directly"""
            return self.encode(text)

    # Create tokenizer instance
    simple_tokenizer = SimpleTokenizer("qwen")

    siliconflow_embedding_available = True

except ImportError:
    print("LightRAG not found. Please install with: pip install lightrag-hku")
    LightRAG = None
    QueryParam = None
    EmbeddingFunc = None
    siliconflow_embed = None
    simple_tokenizer = None
    siliconflow_embedding_available = False

from app.core.config import settings


class KnowledgeSpace:
    """
    Represents a single knowledge space with its own LightRAG instance
    """
    def __init__(self, space_name: str, space_dir: Path, embedding_model: str = "Qwen/Qwen3-Embedding-4B"):
        self.space_name = space_name
        self.space_dir = space_dir
        self.rag = None
        self.documents = []
        self._storages_initialized = False
        self.embedding_model = embedding_model
        self._initialize_rag(force_model=embedding_model)

    def _initialize_rag(self, force_model=None):
        """Initialize LightRAG for this space with optional model override"""
        if not LightRAG:
            print(f"[RAG Init] LightRAG library not available for space '{self.space_name}'. Using fallback mode.")
            self.rag = None
            return

        try:
            print(f"[RAG Init] Initializing LightRAG for knowledge space: {self.space_name}")
            print(f"[RAG Init] Working directory: {self.space_dir}")
            self.space_dir.mkdir(parents=True, exist_ok=True)

            try:
                # Use SiliconFlow async LLM function with retry logic
                async def siliconflow_llm_complete(prompt, model="Pro/moonshotai/Kimi-K2-Instruct-0905", **kwargs):
                    """SiliconFlow async LLM function for LightRAG with configurable model and retry logic"""
                    import httpx

                    max_retries = 3
                    retry_delay = 2  # seconds

                    for attempt in range(max_retries):
                        try:
                            # Get API key from environment
                            api_key = os.getenv("SILICONFLOW_API_KEY", "sk-mczrhquoybaxwicxgsfdaereipgbapmtiersywbzkorjogla")

                            print(f"[LLM] Processing prompt for space '{self.space_name}' with model '{model}', length: {len(prompt)} (attempt {attempt + 1}/{max_retries})")

                            async with httpx.AsyncClient(timeout=180.0) as client:  # Increased timeout to 180s
                                response = await client.post(
                                    "https://api.siliconflow.cn/v1/chat/completions",
                                    headers={
                                        "Authorization": f"Bearer {api_key}",
                                        "Content-Type": "application/json",
                                    },
                                    json={
                                        "model": model,  # Use configurable model
                                        "messages": [{"role": "user", "content": prompt}],
                                        "temperature": 0.1,
                                        "max_tokens": 2048
                                    }
                                )
                                response.raise_for_status()
                                result = response.json()
                                answer = result["choices"][0]["message"]["content"]
                                print(f"[LLM] Successfully generated response for space '{self.space_name}', length: {len(answer)}")
                                return answer
                        except Exception as e:
                            print(f"[LLM] API error for space '{self.space_name}' (attempt {attempt + 1}/{max_retries}): {e}")
                            if attempt < max_retries - 1:
                                print(f"[LLM] Retrying after {retry_delay} seconds...")
                                await asyncio.sleep(retry_delay)
                                retry_delay *= 2  # Exponential backoff
                            else:
                                print(f"[LLM] All retry attempts failed for space '{self.space_name}'")
                                import traceback
                                traceback.print_exc()
                                return f"Error: {str(e)}"

                # Configure embedding model for this space
                embedding_model = force_model or "Qwen/Qwen3-Embedding-4B"  # User's preferred embedding model
                
                # Set the embedding dimension based on the model
                if siliconflow_embedding_available:
                    siliconflow_embed.set_dim(embedding_model)
                    print(f"[RAG Init] Configured embedding model: {embedding_model} with dimension {siliconflow_embed.embedding_dim}")
                    
                    # Check if vector database exists and has dimension mismatch
                    vdb_files_to_check = [
                        "vdb_entities.json",
                        "vdb_relationships.json", 
                        "vdb_chunks.json"
                    ]
                    
                    dimension_mismatch = False
                    current_dim = siliconflow_embed.embedding_dim
                    
                    for vdb_file in vdb_files_to_check:
                        vdb_path = self.space_dir / vdb_file
                        if vdb_path.exists():
                            try:
                                with open(vdb_path, 'r', encoding='utf-8') as f:
                                    vdb_data = json.load(f)
                                    stored_dim = vdb_data.get('embedding_dim', 0)
                                    
                                    if stored_dim != current_dim and stored_dim > 0:
                                        print(f"[RAG Init] Dimension mismatch in {vdb_file}: stored={stored_dim}, current={current_dim}")
                                        dimension_mismatch = True
                                        break
                            except Exception as e:
                                print(f"[RAG Init] Failed to check {vdb_file}: {e}")
                    
                    if dimension_mismatch:
                        print(f"[RAG Init] Clearing vector database to rebuild with correct dimensions...")
                        
                        # Clear all vector database and related files
                        files_to_remove = [
                            "vdb_entities.json",
                            "vdb_relationships.json", 
                            "vdb_chunks.json",
                            "graph_chunk_entity_relation.csv",
                            "graph_chunk_entity_relation.graphml",
                            "kv_store_full_docs.json",
                            "kv_store_text_chunks.json",
                            "kv_store_community_reports.json",
                            "kv_store_llm_response_cache.json",
                            "kv_store_doc_status.json"
                        ]
                        
                        for file_to_remove in files_to_remove:
                            file_path = self.space_dir / file_to_remove
                            if file_path.exists():
                                try:
                                    file_path.unlink()
                                    print(f"[RAG Init] Removed {file_to_remove}")
                                except Exception as e:
                                    print(f"[RAG Init] Failed to remove {file_to_remove}: {e}")
                        
                        print(f"[RAG Init] Vector database cleared. Will rebuild with dimension {current_dim}")
                    else:
                        print(f"[RAG Init] Vector database dimensions are compatible or files don't exist")

                # Initialize LightRAG with custom LLM function
                print(f"[RAG Init] Checking embedding availability: {siliconflow_embedding_available}")

                if siliconflow_embedding_available:
                    print(f"[RAG Init] Creating LightRAG instance with SiliconFlow embeddings...")
                    
                    # Wrap the embedding function in EmbeddingFunc
                    embedding_func_wrapped = EmbeddingFunc(
                        embedding_dim=current_dim,
                        func=siliconflow_embed
                    )
                    
                    self.rag = LightRAG(
                        working_dir=str(self.space_dir),
                        llm_model_func=siliconflow_llm_complete,
                        embedding_func=embedding_func_wrapped,
                        tokenizer=simple_tokenizer  # Use custom simple tokenizer for Qwen
                    )
                    print(f"[RAG Init] SUCCESS: LightRAG instance created successfully for space '{self.space_name}'")
                else:
                    print(f"[RAG Init] Warning: SiliconFlow embedding not available for space '{self.space_name}', LightRAG may not work properly")
                    self.rag = LightRAG(
                        working_dir=str(self.space_dir),
                        llm_model_func=siliconflow_llm_complete,
                        tokenizer=simple_tokenizer  # Use custom simple tokenizer for Qwen
                    )
                    print(f"[RAG Init] SUCCESS: LightRAG instance created (without custom embedding) for space '{self.space_name}'")

                print(f"[RAG Init] SUCCESS: LightRAG initialized successfully for space: {self.space_name}")
                self._load_existing_documents()

            except Exception as e:
                print(f"[RAG Init] ERROR: LightRAG initialization failed for space '{self.space_name}': {e}")
                import traceback
                traceback.print_exc()
                self.rag = None
                print(f"[RAG Init] Using fallback mode for space: {self.space_name}")

        except Exception as e:
            print(f"[RAG Init] ERROR: Failed to initialize LightRAG for space '{self.space_name}': {e}")
            import traceback
            traceback.print_exc()
            self.rag = None

    def reinitialize_with_model(self, embedding_model: str):
        """Reinitialize RAG with a different embedding model"""
        print(f"[RAG Reinit] Reinitializing space '{self.space_name}' with model '{embedding_model}'")
        self.embedding_model = embedding_model
        self._storages_initialized = False
        self._initialize_rag(force_model=embedding_model)
    
    def _load_existing_documents(self):
        """Load existing documents metadata"""
        try:
            docs_file = self.space_dir / "documents.json"
            if docs_file.exists():
                with open(docs_file, 'r', encoding='utf-8') as f:
                    self.documents = json.load(f)
                print(f"Loaded {len(self.documents)} existing documents for space '{self.space_name}'")
        except Exception as e:
            print(f"Failed to load existing documents for space '{self.space_name}': {e}")
            self.documents = []

    def _save_documents(self):
        """Save documents metadata"""
        try:
            docs_file = self.space_dir / "documents.json"
            with open(docs_file, 'w', encoding='utf-8') as f:
                json.dump(self.documents, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Failed to save documents for space '{self.space_name}': {e}")

    async def _ensure_storages_initialized(self):
        """Ensure LightRAG storages and pipeline are initialized"""
        if self.rag and not self._storages_initialized:
            try:
                print(f"[RAG Storage] Initializing storages for space '{self.space_name}'...")
                await self.rag.initialize_storages()
                print(f"[RAG Storage] Initializing pipeline status for space '{self.space_name}'...")
                await initialize_pipeline_status()
                self._storages_initialized = True
                print(f"[RAG Storage] Storages and pipeline initialized successfully for space '{self.space_name}'")
            except Exception as e:
                print(f"[RAG Storage] Failed to initialize storages/pipeline for space '{self.space_name}': {e}")
                raise

    async def add_document(self, content: str, title: str, doc_type: str, metadata=None):
        """Add document to this space"""
        if metadata is None:
            metadata = {}

        lightrag_success = False
        try:
            if self.rag:
                try:
                    # Ensure storages are initialized before first use
                    await self._ensure_storages_initialized()

                    print(f"[Document] Adding document '{title}' to space '{self.space_name}', content length: {len(content)}")

                    rag_metadata = {
                        "title": title,
                        "type": doc_type,
                        "timestamp": str(Path(settings.UPLOAD_DIR).stat().st_mtime)
                    }
                    rag_metadata.update(metadata)

                    print(f"[Document] Calling LightRAG ainsert for '{title}'...")
                    await self.rag.ainsert(content)

                    lightrag_success = True
                    print(f"[Document] Successfully added document to LightRAG for space '{self.space_name}': {title}")
                except Exception as e:
                    print(f"[Document] LightRAG insertion failed for space '{self.space_name}': {e}")
                    import traceback
                    traceback.print_exc()
                    # Don't return False yet, we'll add metadata with error status
            else:
                print(f"[Document] LightRAG not available for space '{self.space_name}', skipping indexing")

            # Add document metadata with status
            doc_info = {
                "title": title,
                "type": doc_type,
                "content_length": len(content),
                "timestamp": str(Path(settings.UPLOAD_DIR).stat().st_mtime),
                "status": "active" if lightrag_success or not self.rag else "metadata_only",
                "knowledge_space": self.space_name,
                "lightrag_indexed": lightrag_success
            }
            doc_info.update(metadata)
            self.documents.append(doc_info)
            self._save_documents()

            print(f"[Document] Document metadata saved for '{title}', indexed: {lightrag_success}")
            return True

        except Exception as e:
            print(f"[Document] Failed to add document to space '{self.space_name}': {e}")
            import traceback
            traceback.print_exc()
            return False

    async def delete_document(self, document_title: str) -> bool:
        """Delete document from this space"""
        try:
            original_len = len(self.documents)
            self.documents = [doc for doc in self.documents if doc.get("title") != document_title]

            if len(self.documents) == original_len:
                print(f"Document '{document_title}' not found in space '{self.space_name}'")
                return False

            self._save_documents()
            print(f"Removed document '{document_title}' from space '{self.space_name}'")
            return True

        except Exception as e:
            print(f"Failed to delete document '{document_title}' from space '{self.space_name}': {e}")
            return False

    async def query_knowledge(self, question: str, mode: str = "hybrid", model: Optional[str] = None):
        """Query knowledge in this space with optional model override"""
        if not self.rag:
            return {"answer": f"Knowledge base (LightRAG) not available for space '{self.space_name}'", "sources": [], "mode": mode}

        try:
            import time
            start_time = time.time()
            print(f"[Query] Starting query for space '{self.space_name}': '{question[:50]}...'")
            
            # Ensure storages are initialized before querying
            init_start = time.time()
            await self._ensure_storages_initialized()
            init_time = time.time() - init_start
            print(f"[Query] Storage initialization took {init_time:.2f} seconds")

            # Create query params with optional model override and disable rerank
            query_param = QueryParam(mode=mode, enable_rerank=False)
            
            # Use user-specified model or default to Pro/moonshotai/Kimi-K2-Instruct-0905
            selected_model = model or "Pro/moonshotai/Kimi-K2-Instruct-0905"
            print(f"[Query] Using model: {selected_model}")
            
            # Create a custom LLM function for this query with the specified model
            async def query_specific_llm_complete(prompt, **kwargs):
                """SiliconFlow async LLM function with query-specific model"""
                import httpx
                import asyncio

                max_retries = 3
                retry_delay = 2  # seconds

                for attempt in range(max_retries):
                    try:
                        # Get API key from environment
                        api_key = os.getenv("SILICONFLOW_API_KEY", "sk-mczrhquoybaxwicxgsfdaereipgbapmtiersywbzkorjogla")

                        print(f"[LLM] Processing prompt for space '{self.space_name}' with model '{selected_model}', length: {len(prompt)} (attempt {attempt + 1}/{max_retries})")

                        async with httpx.AsyncClient(timeout=180.0) as client:  # Increased timeout to 180s
                            response = await client.post(
                                "https://api.siliconflow.cn/v1/chat/completions",
                                headers={
                                    "Authorization": f"Bearer {api_key}",
                                    "Content-Type": "application/json",
                                },
                                json={
                                    "model": selected_model,  # Use query-specific model
                                    "messages": [{"role": "user", "content": prompt}],
                                    "temperature": 0.1,
                                    "max_tokens": 2048
                                }
                            )
                            response.raise_for_status()
                            result = response.json()
                            answer = result["choices"][0]["message"]["content"]
                            print(f"[LLM] Successfully generated response for space '{self.space_name}', length: {len(answer)}")
                            return answer
                    except Exception as e:
                        print(f"[LLM] API error for space '{self.space_name}' (attempt {attempt + 1}/{max_retries}): {e}")
                        if attempt < max_retries - 1:
                            print(f"[LLM] Retrying after {retry_delay} seconds...")
                            await asyncio.sleep(retry_delay)
                            retry_delay *= 2  # Exponential backoff
                        else:
                            print(f"[LLM] All retry attempts failed for space '{self.space_name}'")
                            import traceback
                            traceback.print_exc()
                            return f"Error: {str(e)}"
            
            # Temporarily replace the RAG's LLM function with our query-specific one
            original_llm_func = self.rag.llm_model_func
            self.rag.llm_model_func = query_specific_llm_complete
            
            try:
                # Perform the actual query with the custom model
                query_start = time.time()
                result = await self.rag.aquery(question, param=query_param)
                query_time = time.time() - query_start
                
                total_time = time.time() - start_time
                print(f"[Query] LightRAG query took {query_time:.2f} seconds")
                print(f"[Query] Total query time: {total_time:.2f} seconds")
                
                return {
                    "answer": str(result),
                    "sources": [],
                    "mode": mode,
                    "model_used": selected_model,
                    "query_time": total_time
                }
            finally:
                # Restore the original LLM function
                self.rag.llm_model_func = original_llm_func
                
        except Exception as e:
            print(f"Failed to query knowledge base for space '{self.space_name}': {e}")
            import traceback
            traceback.print_exc()
            return {"answer": f"Error: {str(e)}", "sources": [], "mode": mode}

    def get_documents_list(self):
        """Get documents list for this space"""
        return self.documents.copy()

    def is_available(self) -> bool:
        """Check if knowledge base is available for this space"""
        return self.rag is not None

    def get_stats(self):
        """Get statistics for this space"""
        stats = {
            "space_name": self.space_name,
            "lightrag_available": self.is_available(),
            "total_documents": len(self.documents),
            "document_types": {},
            "total_content_length": 0
        }

        for doc in self.documents:
            doc_type = doc.get("type", "unknown")
            stats["document_types"][doc_type] = stats["document_types"].get(doc_type, 0) + 1
            stats["total_content_length"] += doc.get("content_length", 0)

        return stats

    async def process_file_upload(self, file_content: bytes, filename: str,
                                title: Optional[str] = None, user_id: Optional[str] = None,
                                course_name: Optional[str] = None,
                                space_name: Optional[str] = None,
                                model: Optional[str] = None) -> Dict[str, Any]:
        """Process uploaded file for this space"""
        try:
            import mimetypes

            # Determine file type
            content_type, _ = mimetypes.guess_type(filename)
            file_ext = Path(filename).suffix.lower()

            print(f"[FileProcess] Processing file: {filename}, type: {content_type}, ext: {file_ext}, size: {len(file_content)} bytes")

            text_content = ""
            doc_type = "unknown"

            # Extract text based on file type
            if file_ext == '.txt':
                print(f"[FileProcess] Extracting plain text from {filename}")
                try:
                    text_content = file_content.decode('utf-8', errors='ignore')
                    doc_type = "text/plain"
                    print(f"[FileProcess] Extracted {len(text_content)} characters from TXT")
                except Exception as e:
                    return {"success": False, "error": f"Failed to decode text file: {str(e)}"}

            elif file_ext == '.md':
                print(f"[FileProcess] Extracting markdown from {filename}")
                try:
                    text_content = file_content.decode('utf-8', errors='ignore')
                    doc_type = "text/markdown"
                    print(f"[FileProcess] Extracted {len(text_content)} characters from Markdown")
                except Exception as e:
                    return {"success": False, "error": f"Failed to decode markdown file: {str(e)}"}

            elif file_ext == '.pdf':
                print(f"[FileProcess] Extracting text from PDF: {filename}")
                try:
                    # Try PyMuPDF (fitz) first
                    import fitz
                    doc = fitz.open(stream=file_content, filetype="pdf")
                    text_content = ""
                    page_count = len(doc)
                    print(f"[FileProcess] PDF has {page_count} pages")

                    for page_num, page in enumerate(doc, 1):
                        page_text = page.get_text()
                        text_content += page_text
                        print(f"[FileProcess] Extracted {len(page_text)} chars from page {page_num}/{page_count}")

                    doc.close()
                    doc_type = "application/pdf"
                    print(f"[FileProcess] Total extracted: {len(text_content)} characters from PDF")

                except ImportError:
                    print("[FileProcess] PyMuPDF not available, trying pdfplumber...")
                    try:
                        import pdfplumber
                        import io
                        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                            text_content = ""
                            for page_num, page in enumerate(pdf.pages, 1):
                                page_text = page.extract_text() or ""
                                text_content += page_text
                                print(f"[FileProcess] Extracted {len(page_text)} chars from page {page_num}/{len(pdf.pages)}")
                        doc_type = "application/pdf"
                        print(f"[FileProcess] Total extracted: {len(text_content)} characters from PDF")
                    except ImportError:
                        return {
                            "success": False,
                            "error": "PDF processing requires PyMuPDF (fitz) or pdfplumber. Install with: pip install PyMuPDF"
                        }
                    except Exception as e:
                        print(f"[FileProcess] PDF parsing failed: {e}")
                        import traceback
                        traceback.print_exc()
                        return {"success": False, "error": f"PDF parsing failed: {str(e)}"}
                except Exception as e:
                    print(f"[FileProcess] PDF parsing failed: {e}")
                    import traceback
                    traceback.print_exc()
                    return {"success": False, "error": f"PDF parsing failed: {str(e)}"}

            elif file_ext in ['.docx', '.doc']:
                print(f"[FileProcess] Extracting text from DOCX: {filename}")
                try:
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(file_content))

                    # Extract paragraphs
                    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
                    text_content = "\n".join(paragraphs)

                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join([cell.text.strip() for cell in row.cells])
                            if row_text.strip():
                                text_content += "\n" + row_text

                    doc_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    print(f"[FileProcess] Extracted {len(text_content)} characters from DOCX")

                except ImportError:
                    return {
                        "success": False,
                        "error": "DOCX processing requires python-docx. Install with: pip install python-docx"
                    }
                except Exception as e:
                    print(f"[FileProcess] DOCX parsing failed: {e}")
                    import traceback
                    traceback.print_exc()
                    return {"success": False, "error": f"DOCX parsing failed: {str(e)}"}

            else:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_ext}. Supported: .txt, .md, .pdf, .docx"
                }

            # Validate extracted content
            if not text_content.strip():
                print(f"[FileProcess] Warning: No readable content extracted from {filename}")
                return {
                    "success": False,
                    "error": "No readable content found in file. The file may be empty or corrupted."
                }

            # Prepare metadata
            metadata = {
                "filename": filename,
                "user_id": user_id,
                "file_size": len(file_content),
                "content_length": len(text_content),
                "file_ext": file_ext
            }
            if course_name:
                metadata["course_name"] = course_name
            if space_name:
                metadata["space_name"] = space_name

            doc_title = title or Path(filename).stem

            print(f"[FileProcess] Adding document '{doc_title}' to knowledge base...")

            # Add to knowledge base
            success = await self.add_document(
                content=text_content,
                title=doc_title,
                doc_type=doc_type,
                metadata=metadata
            )

            if success:
                return {
                    "success": True,
                    "message": f"Document '{doc_title}' added to knowledge space '{self.space_name}' successfully",
                    "filename": filename,
                    "doc_type": doc_type,
                    "content_length": len(text_content),
                    "file_size": len(file_content),
                    "lightrag_available": self.rag is not None
                }
            else:
                return {
                    "success": False,
                    "error": "Failed to add document to knowledge base"
                }

        except Exception as e:
            print(f"[FileProcess] File processing failed for {filename}: {e}")
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": f"File processing failed: {str(e)}"
            }

    def get_graph(self) -> Dict[str, Any]:
        """Get knowledge graph data from this space"""
        # This would be a LightRAG visualization method if available
        # For now, return mock data
        return {
            "entities": [
                {
                    "id": "1",
                    "label": "Sample Concept",
                    "type": "concept",
                    "properties": {}
                }
            ],
            "relationships": [],
            "chunks": []
        }


class KnowledgeBaseService:
    """
    Multi-space knowledge base service manager
    """

    def __init__(self):
        """Initialize the knowledge base service"""
        self.spaces = {}
        self.default_space = "default"
        self._load_existing_spaces()
        self._ensure_default_space()

    def _load_existing_spaces(self):
        """Load existing knowledge spaces"""
        try:
            kb_root_dir = Path(settings.UPLOAD_DIR) / "knowledge_spaces"
            kb_root_dir.mkdir(parents=True, exist_ok=True)

            # Load all spaces from knowledge_spaces directory
            for space_dir in kb_root_dir.iterdir():
                if space_dir.is_dir():
                    space_name = space_dir.name
                    # Use the user's preferred embedding model
                    space = KnowledgeSpace(space_name, space_dir, "Qwen/Qwen3-Embedding-4B")
                    self.spaces[space_name] = space
                    print(f"Loaded space: {space_name}")

            # Migrate legacy knowledge_base directory to default space if needed
            old_kb_dir = Path(settings.UPLOAD_DIR) / "knowledge_base"
            if old_kb_dir.exists() and old_kb_dir.is_dir():
                # Only migrate if default doesn't exist yet
                if "default" not in self.spaces:
                    print("Found legacy knowledge_base directory, migrating to default space...")
                    default_dir = kb_root_dir / "default"
                    default_dir.mkdir(parents=True, exist_ok=True)

                    # Copy documents.json if it exists
                    old_docs_file = old_kb_dir / "documents.json"
                    if old_docs_file.exists():
                        import shutil
                        new_docs_file = default_dir / "documents.json"
                        shutil.copy2(old_docs_file, new_docs_file)
                        print(f"Migrated documents.json from legacy directory")

                    # Load the default space from new location
                    default_space = KnowledgeSpace("default", default_dir, "Qwen/Qwen3-Embedding-4B")
                    self.spaces["default"] = default_space
                    print("Legacy knowledge base migrated to default space")
                else:
                    print("Default space already exists, skipping legacy migration")

            print(f"Loaded {len(self.spaces)} knowledge space(s)")

        except Exception as e:
            print(f"Failed to load existing spaces: {e}")
            self.spaces = {}

    def _ensure_default_space(self):
        """Ensure the default space exists"""
        if "default" not in self.spaces:
            kb_root_dir = Path(settings.UPLOAD_DIR) / "knowledge_spaces"
            default_dir = kb_root_dir / "default"
            default_space = KnowledgeSpace("default", default_dir, "Qwen/Qwen3-Embedding-4B")
            self.spaces["default"] = default_space
            print("Created default knowledge space")

    def create_space(self, space_name: str) -> bool:
        """Create a new knowledge space"""
        try:
            if space_name in self.spaces:
                print(f"Space '{space_name}' already exists")
                return False

            invalid_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|', '\0']
            if any(char in space_name for char in invalid_chars):
                print(f"Invalid space name '{space_name}': contains forbidden characters")
                return False

            if not space_name.strip():
                print("Space name cannot be empty")
                return False

            kb_root_dir = Path(settings.UPLOAD_DIR) / "knowledge_spaces"
            space_dir = kb_root_dir / space_name

            space = KnowledgeSpace(space_name, space_dir, "Qwen/Qwen3-Embedding-4B")
            self.spaces[space_name] = space
            print(f"Successfully created knowledge space: {space_name}")
            return True

        except Exception as e:
            print(f"Failed to create space '{space_name}': {e}")
            return False

    def get_space(self, space_name: Optional[str] = None) -> Optional[KnowledgeSpace]:
        """Get a knowledge space by name, defaults to default space"""
        space_name = space_name or self.default_space
        return self.spaces.get(space_name)

    def list_spaces(self) -> List[str]:
        """Get list of all space names"""
        return list(self.spaces.keys())

    def get_spaces_info(self) -> List[Dict[str, Any]]:
        """Get detailed information about all spaces"""
        spaces_info = []
        for space_name, space in self.spaces.items():
            spaces_info.append(space.get_stats())
        return spaces_info

    async def add_document(self, content: str, title: str, doc_type: str,
                          metadata: Optional[Dict] = None, space_name: Optional[str] = None) -> bool:
        """Add document to specific space (defaults to default)"""
        space = self.get_space(space_name)
        if not space:
            print(f"Space '{space_name}' not found")
            return False
        return await space.add_document(content, title, doc_type, metadata)

    async def delete_document(self, document_title: str, space_name: Optional[str] = None) -> bool:
        """Delete document from specific space (defaults to default)"""
        space = self.get_space(space_name)
        if not space:
            print(f"Space '{space_name}' not found")
            return False
        return await space.delete_document(document_title)

    async def query_knowledge(self, question: str, mode: str = "hybrid",
                             space_name: Optional[str] = None, model: Optional[str] = None) -> Dict[str, Any]:
        """Query specific knowledge space with optional model override (defaults to default)"""
        space = self.get_space(space_name)
        if not space:
            return {"answer": f"Knowledge space '{space_name}' not found", "sources": [], "mode": mode}
        return await space.query_knowledge(question, mode, model)

    def get_documents_list(self, space_name: Optional[str] = None) -> List[Dict[str, Any]]:
        """Get documents list for specific space (defaults to default)"""
        space = self.get_space(space_name)
        if not space:
            return []
        return space.get_documents_list()

    def is_available(self, space_name: Optional[str] = None) -> bool:
        """Check if specific knowledge space is available (defaults to default)"""
        space = self.get_space(space_name)
        return space.is_available() if space else False

    def get_stats(self, space_name: Optional[str] = None) -> Dict[str, Any]:
        """Get statistics for specific space (defaults to default)"""
        space = self.get_space(space_name)
        if not space:
            return {
                "space_name": space_name or "unknown",
                "lightrag_available": False,
                "total_documents": 0,
                "document_types": {},
                "total_content_length": 0
            }
        return space.get_stats()

    async def process_file_upload(self, file_content: bytes, filename: str,
                                title: Optional[str] = None, user_id: Optional[str] = None,
                                course_name: Optional[str] = None,
                                space_name: Optional[str] = None,
                                model: Optional[str] = None) -> Dict[str, Any]:
        """Process file upload for specific space (defaults to default)"""
        space = self.get_space(space_name)
        if not space:
            return {"success": False, "error": f"Knowledge space '{space_name}' not found"}

        return await space.process_file_upload(
            file_content=file_content,
            filename=filename,
            title=title,
            user_id=user_id,
            course_name=course_name,
            space_name=space_name,
            model=model
        )

    async def get_course_structure(self, course_name: str, space_name: Optional[str] = None) -> Dict[str, Any]:
        """Get course structure for specific space"""
        space = self.get_space(space_name)
        if not space:
            return {"course_name": course_name, "modules": [], "error": "Space not found"}
        prompt = f"Analyze the course \"{course_name}\" and provide: 1. Main modules/topics 2. Key knowledge points for each module 3. Difficulty assessment. Return structured JSON."
        result = await space.query_knowledge(prompt, mode="global")
        return {
            "course_name": course_name,
            "modules": [],
            "total_knowledge_points": 0,
            "analysis": result["answer"]
        }

    async def analyze_assignment_errors(self, assignment_content: str, course_name: str, space_name: Optional[str] = None) -> Dict[str, Any]:
        """Analyze assignment for specific space"""
        space = self.get_space(space_name)
        if not space:
            return {"analysis": "Space not found", "error": "knowledge space not found"}
        prompt = f"Analyze this assignment submission for the course \"{course_name}\" and identify: 1. Understanding points 2. Difficulties 3. Error patterns 4. Suggestions."
        result = await space.query_knowledge(prompt, mode="hybrid")
        return {
            "analysis": result["answer"],
            "knowledge_gaps": [],
            "error_patterns": [],
            "suggestions": []
        }


# Create global knowledge base service instance
knowledge_base = KnowledgeBaseService()