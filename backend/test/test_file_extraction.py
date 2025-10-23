"""
测试文件提取功能的脚本
运行方式: python test_file_extraction.py
"""
import asyncio
import sys
from pathlib import Path

# 添加app目录到路径
sys.path.insert(0, str(Path(__file__).parent))

async def test_file_extraction():
    """测试各种文件格式的提取"""
    print("=" * 60)
    print("测试文件提取功能")
    print("=" * 60)

    from app.services.knowledge_base import KnowledgeSpace

    # 创建临时知识空间
    test_dir = Path("./test_uploads")
    test_dir.mkdir(exist_ok=True)

    space = KnowledgeSpace("test", test_dir)

    # 测试1: TXT文件
    print("\n1. 测试TXT文件提取...")
    txt_content = "这是一个测试文本文件。\n包含中文和English混合内容。".encode('utf-8')
    result = await space.process_file_upload(
        file_content=txt_content,
        filename="test.txt",
        title="测试文本"
    )
    print(f"   结果: {'✅ 成功' if result['success'] else '❌ 失败'}")
    if result['success']:
        print(f"   提取字符数: {result['content_length']}")
    else:
        print(f"   错误: {result.get('error')}")

    # 测试2: Markdown文件
    print("\n2. 测试Markdown文件提取...")
    md_content = """# 标题
## 子标题
这是一段**加粗**的文字。
- 列表项1
- 列表项2
""".encode('utf-8')
    result = await space.process_file_upload(
        file_content=md_content,
        filename="test.md",
        title="测试Markdown"
    )
    print(f"   结果: {'✅ 成功' if result['success'] else '❌ 失败'}")
    if result['success']:
        print(f"   提取字符数: {result['content_length']}")
    else:
        print(f"   错误: {result.get('error')}")

    # 测试3: PDF文件
    print("\n3. 测试PDF文件提取...")
    print("   检查PyMuPDF...")
    try:
        import fitz
        print("   ✅ PyMuPDF已安装")

        # 创建一个简单的PDF用于测试
        try:
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((100, 100), "这是一个测试PDF文件。\n包含中文内容。", fontsize=12)
            pdf_bytes = doc.tobytes()
            doc.close()

            result = await space.process_file_upload(
                file_content=pdf_bytes,
                filename="test.pdf",
                title="测试PDF"
            )
            print(f"   结果: {'✅ 成功' if result['success'] else '❌ 失败'}")
            if result['success']:
                print(f"   提取字符数: {result['content_length']}")
            else:
                print(f"   错误: {result.get('error')}")
        except Exception as e:
            print(f"   ⚠️  创建测试PDF失败: {e}")
            print(f"   提示: 请手动上传PDF文件测试")

    except ImportError:
        print("   ❌ PyMuPDF未安装")
        print("   安装命令: pip install PyMuPDF")

    # 测试4: DOCX文件
    print("\n4. 测试DOCX文件提取...")
    print("   检查python-docx...")
    try:
        from docx import Document
        import io
        print("   ✅ python-docx已安装")

        # 创建一个简单的DOCX用于测试
        try:
            doc = Document()
            doc.add_heading('测试标题', 0)
            doc.add_paragraph('这是第一段文字。')
            doc.add_paragraph('这是第二段文字。')

            # 添加表格
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = '列1'
            table.rows[0].cells[1].text = '列2'
            table.rows[1].cells[0].text = '数据1'
            table.rows[1].cells[1].text = '数据2'

            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_bytes = docx_io.getvalue()

            result = await space.process_file_upload(
                file_content=docx_bytes,
                filename="test.docx",
                title="测试DOCX"
            )
            print(f"   结果: {'✅ 成功' if result['success'] else '❌ 失败'}")
            if result['success']:
                print(f"   提取字符数: {result['content_length']}")
                print(f"   包含表格: 是")
            else:
                print(f"   错误: {result.get('error')}")
        except Exception as e:
            print(f"   ⚠️  创建测试DOCX失败: {e}")
            print(f"   提示: 请手动上传DOCX文件测试")

    except ImportError:
        print("   ❌ python-docx未安装")
        print("   安装命令: pip install python-docx")

    # 测试5: 不支持的文件类型
    print("\n5. 测试不支持的文件类型...")
    result = await space.process_file_upload(
        file_content=b"test data",
        filename="test.xyz",
        title="不支持的文件"
    )
    print(f"   结果: {'✅ 正确拒绝' if not result['success'] else '❌ 错误接受'}")
    if not result['success']:
        print(f"   错误信息: {result.get('error')}")

    # 测试6: 空文件
    print("\n6. 测试空文件...")
    result = await space.process_file_upload(
        file_content=b"",
        filename="empty.txt",
        title="空文件"
    )
    print(f"   结果: {'✅ 正确拒绝' if not result['success'] else '❌ 错误接受'}")
    if not result['success']:
        print(f"   错误信息: {result.get('error')}")

    print("\n" + "=" * 60)
    print("测试完成!")
    print("=" * 60)

    # 清理
    print("\n清理测试文件...")
    import shutil
    if test_dir.exists():
        shutil.rmtree(test_dir)
    print("清理完成!")

if __name__ == "__main__":
    print("开始测试文件提取功能...\n")

    # 检查依赖
    print("检查依赖包...")
    dependencies = {
        "PyMuPDF (fitz)": False,
        "python-docx": False,
        "pdfplumber": False
    }

    try:
        import fitz
        dependencies["PyMuPDF (fitz)"] = True
    except ImportError:
        pass

    try:
        from docx import Document
        dependencies["python-docx"] = True
    except ImportError:
        pass

    try:
        import pdfplumber
        dependencies["pdfplumber"] = True
    except ImportError:
        pass

    print("\n依赖包状态:")
    for pkg, installed in dependencies.items():
        status = "✅ 已安装" if installed else "❌ 未安装"
        print(f"  {pkg}: {status}")

    if not dependencies["PyMuPDF (fitz)"] and not dependencies["pdfplumber"]:
        print("\n⚠️  警告: PDF处理库未安装")
        print("   推荐安装: pip install PyMuPDF")

    if not dependencies["python-docx"]:
        print("\n⚠️  警告: DOCX处理库未安装")
        print("   推荐安装: pip install python-docx")

    print()

    # 运行测试
    asyncio.run(test_file_extraction())
